r"""
csv_word_merge.py
-John Taylor
May-3-2021

Merge CSV fields into a MS Word template

References: https://stackoverflow.com/a/61516850/452281
"""

import argparse
import csv
import os
import time
from docx import Document
from docx2pdf import convert

pgm_name = "csv_word_merge"
pgm_version = "1.0.0"
pgm_url = "https://github.com/jftuga/csv_word_merge"

def get_csv_data(fname: str) -> dict:
    all_rows = []
    with open(fname, newline="") as c:
        reader = csv.DictReader(c)
        for row in reader:
            all_rows.append(row)
    return all_rows


def get_docx_name(row: dict, output: int) -> str:
    return "%s.docx" % row[output]


# https://stackoverflow.com/a/55733040/452281
def docx_replace(doc, data) -> int:
    count = 0
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in doc.paragraphs:
        for key, val in data.items():
            key = key.strip()
            val = val.strip()
            key_name = f"_{key}_"
            if key_name in p.text:
                count += 1
                inline = p.runs
                # Replace strings and retain the same style.
                # The text to be replaced can be split over several runs so
                # search through, identify which runs need to have text replaced
                # then replace the text in those identified
                started = False
                key_index = 0
                # found_runs is a list of (inline index, index of match, length of match)
                found_runs = list()
                found_all = False
                replace_done = False
                for i in range(len(inline)):

                    # case 1: found in single run so short circuit the replace
                    if key_name in inline[i].text and not started:
                        found_runs.append((i, inline[i].text.find(key_name), len(key_name)))
                        text = inline[i].text.replace(key_name, str(val))
                        inline[i].text = text
                        replace_done = True
                        found_all = True
                        break

                    if key_name[key_index] not in inline[i].text and not started:
                        # keep looking ...
                        continue

                    # case 2: search for partial text, find first run
                    if key_name[key_index] in inline[i].text and inline[i].text[-1] in key_name and not started:
                        # check sequence
                        start_index = inline[i].text.find(key_name[key_index])
                        check_length = len(inline[i].text)
                        for text_index in range(start_index, check_length):
                            if inline[i].text[text_index] != key_name[key_index]:
                                # no match so must be false positive
                                break
                        if key_index == 0:
                            started = True
                        chars_found = check_length - start_index
                        key_index += chars_found
                        found_runs.append((i, start_index, chars_found))
                        if key_index != len(key_name):
                            continue
                        else:
                            # found all chars in key_name
                            found_all = True
                            break

                    # case 2: search for partial text, find subsequent run
                    if key_name[key_index] in inline[i].text and started and not found_all:
                        # check sequence
                        chars_found = 0
                        check_length = len(inline[i].text)
                        for text_index in range(0, check_length):
                            if inline[i].text[text_index] == key_name[key_index]:
                                key_index += 1
                                chars_found += 1
                            else:
                                break
                        # no match so must be end
                        found_runs.append((i, 0, chars_found))
                        if key_index == len(key_name):
                            found_all = True
                            break

                if found_all and not replace_done:
                    for i, item in enumerate(found_runs):
                        index, start, length = [t for t in item]
                        if i == 0:
                            text = inline[index].text.replace(inline[index].text[start:start + length], str(val))
                            inline[index].text = text
                        else:
                            text = inline[index].text.replace(inline[index].text[start:start + length], '')
                            inline[index].text = text
    return count


def create_dest(dest: str):
    if not os.path.exists(dest):
        os.mkdir(dest, 0o755)


def main():
    version_string = f"{pgm_name}, v{pgm_version}, {pgm_url}"

    parser = argparse.ArgumentParser(description="Word Macro Automation")
    parser.add_argument("--csv", "-c", help="csv file containing macros", required=True)
    parser.add_argument("--col", "-C", help="column name for output PDF", required=True)
    parser.add_argument("--dest", "-d", help="destination folder", required=True)
    parser.add_argument("--version", "-v", help="display version and then exit", action="version", version=version_string)
    parser.add_argument("wordfile", metavar="wordfile", help="MS Word file with macros")
    args = parser.parse_args()

    create_dest(args.dest)
    csvdata = get_csv_data(args.csv)
    for row in csvdata:
        document = Document(args.wordfile)
        changes = 0
        if len(row[args.col]) == 0:
            print("Skipping row, invalid data: ", row)
            continue
        changes += docx_replace(document, row)

        print(f"{row=}; number of changes: {changes}")
        if changes > 0:
            docx_name = get_docx_name(row, args.col)
            docx_name = os.path.join(args.dest, docx_name)
            document.save(docx_name)
            time.sleep(0.1)
            print(f"{docx_name=}")
            convert(docx_name)
        else:
            print(f"No changes found for {row}")


if "__main__" == __name__:
    main()

# end of script
